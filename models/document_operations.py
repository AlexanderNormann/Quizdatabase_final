import os

from docx import Document
from datetime import datetime

from models.database import execute_query, fetch_rows


def create_document(questions):
    document = Document()

    for question in questions:
        question_text = question[1]
        answer_text = question[2]

        document.add_paragraph(question_text)
        document.add_paragraph("Answer: " + answer_text)
        document.add_paragraph("")

    return document


def save_document(document):
    desktop_dir = os.path.join(os.path.expanduser("~"), "Quizzer")

    folder_name = "quiz_documents"
    folder_path = os.path.join(desktop_dir, folder_name)

    if not os.path.exists(folder_path):
        os.makedirs(folder_path)

    current_date = datetime.now().strftime("%Y-%m-%d")

    document_filename = f"quiz_document_{current_date}.docx"

    document_path = os.path.join(folder_path, document_filename)

    document.save(document_path)
    return document_path


def update_last_used(question_ids):
    updated_last_used = []

    for question_id in question_ids:
        update_question_last_used(question_id)
        last_used = get_question_last_used(question_id)
        updated_last_used.append(last_used)

    return updated_last_used


def filter_rows_by_date(rows, start_date, end_date):
    filtered_rows = []

    for row in rows:
        row_last_used = row[4]
        if start_date and end_date:
            if start_date <= row_last_used <= end_date:
                filtered_rows.append(row)
        elif start_date:
            if start_date <= row_last_used:
                filtered_rows.append(row)
        elif end_date:
            if row_last_used <= end_date:
                filtered_rows.append(row)
        else:
            filtered_rows.append(row)

    return filtered_rows


def update_question_last_used(question_id):
    current_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    execute_query("UPDATE questions SET last_used=? WHERE id=?", current_time, question_id)


def get_question_last_used(question_id):
    rows = fetch_rows("SELECT last_used FROM questions WHERE id=?", question_id)
    last_used = rows[0][0] if rows else None
    return last_used
